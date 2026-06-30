import os
import io
import pandas as pd
from docx import Document
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from services.rag import query_sla, sla_is_ingested

load_dotenv()

class DocxHeader(BaseModel):
    vendor_name: str = Field(description="The company providing the service. Do NOT output the airport name.")
    metric_or_service: str = Field(description="The specific service or metric being measured.")

class VendorClause(BaseModel):
    reasoning: str = Field(description="Step-by-step reasoning. 1. Read the retrieved pages. 2. Find the section for the TARGET VENDOR. 3. Identify their specific table on that page. 4. Extract threshold and penalty. DO NOT look at tables belonging to other vendors.")
    vendor_name: str = Field(description="Full legal vendor name exactly as written.")
    metric: str = Field(description="The exact metric name from the TARGET VENDOR'S table. Do not copy metrics from other vendors.")
    threshold_value: float = Field(description="The limit number. Return 0 if the target vendor is missing.")
    threshold_unit: str = Field(description="Time unit: 'seconds', 'minutes', 'hours'.")
    penalty_per_breach: int = Field(description="Penalty as plain integer. Return 0 if missing.")
    currency: str = Field(description="Always INR")

def to_minutes(value: float, unit: str) -> float:
    if value == 0: return 0
    unit = (unit or "").strip().lower()
    if unit.startswith("second"): return value / 60
    if unit.startswith("hour"): return value * 60
    return value

def extract_tables_from_docx(file_bytes: bytes) -> list:
    dfs = []
    for table in Document(io.BytesIO(file_bytes)).tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) > 1: dfs.append(pd.DataFrame(rows[1:], columns=rows[0]))
    return dfs

def extract_text_from_docx(file_bytes: bytes) -> str:
    return "\n".join([p.text.strip() for p in Document(io.BytesIO(file_bytes)).paragraphs if p.text.strip()])

def get_sla_thresholds(doc_meta_text: str) -> dict:
    if not sla_is_ingested(): 
        return {"extraction_failed": True, "error": "No SLA ingested. Please upload an SLA PDF via /ingest/sla first."}

    llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0)
    
    hint_prompt = ChatPromptTemplate.from_template(
        "Extract the Vendor Name (the company providing the service, NOT the airport) and the Service from this header:\n\n{header}"
    )
    try:
        hints = (hint_prompt | llm.with_structured_output(DocxHeader)).invoke({"header": doc_meta_text[:800]})
        vendor_hint = hints.vendor_name
        metric_hint = hints.metric_or_service
    except Exception as e:
        return {"extraction_failed": True, "error": f"Failed to parse DOCX header: {e}"}

    rag_query = f"{vendor_hint} {metric_hint} maximum permissible duration threshold penalty"
    sla_context = query_sla(rag_query, k=4) # Fetch top 4 pages

    clause_prompt = ChatPromptTemplate.from_template("""
    You are an elite legal contract AI parsing FAISS vector data.
    
    PAGES RETRIEVED FROM FAISS: 
    {sla_context}
    
    TARGET VENDOR: {vendor_hint} 
    TARGET METRIC/SERVICE: {metric_hint}
    
    CRITICAL RULES:
    1. Populate the 'reasoning' field first. You will see entire pages. You MUST locate the text for {vendor_hint} and only extract data from the table immediately associated with them.
    2. Ignore tables belonging to other vendors (e.g. Vega Ramp Solutions, Petrosafe) even if they are on the same page.
    3. If {vendor_hint} is not found on these pages, return 0 for threshold_value and penalty_per_breach.
    """)
    
    try:
        res: VendorClause = (clause_prompt | llm.with_structured_output(VendorClause)).invoke({
            "sla_context": sla_context, "vendor_hint": vendor_hint, "metric_hint": metric_hint
        })
        
        if res.threshold_value == 0 or res.penalty_per_breach == 0:
            return {
                "extraction_failed": True, 
                "error": f"SLA rules for '{vendor_hint}' not found in the FAISS database. Please upload the correct SLA PDF for this vendor."
            }

        return {
            "vendor": res.vendor_name, "metric": res.metric, "currency": res.currency,
            "threshold_minutes": to_minutes(res.threshold_value, res.threshold_unit),
            "penalty_per_breach": res.penalty_per_breach, "extraction_failed": False
        }
    except Exception as e:
        return {"extraction_failed": True, "error": str(e)}

def detect_breaches_from_df(df: pd.DataFrame, limit: float, penalty: int, vendor: str, metric: str) -> list:
    breaches = []
    dur_col = next((c for c in df.columns if any(s in c.lower() for s in ["dur", "elapsed"])), None)
    flt_col = next((c for c in df.columns if any(s in c.lower() for s in ["flight", "flt"])), None)
    dat_col = next((c for c in df.columns if "date" in c.lower() or "day" in c.lower()), None)

    if not dur_col: return breaches
    for _, row in df.iterrows():
        try: duration = float(str(row[dur_col]).replace("min", "").strip())
        except: continue
        
        # STRICT LOGIC: Breach ONLY occurs if duration strictly exceeds the limit
        if duration > limit:
            breaches.append({
                "vendor": vendor, "metric": metric, "penalty": penalty,
                "flight_number": str(row[flt_col]).strip() if flt_col else "N/A",
                "date": str(row[dat_col]).strip() if dat_col else "N/A",
                "actual_minutes": round(duration, 1), "threshold_minutes": limit,
                "overage_minutes": round(duration - limit, 1),
            })
    return breaches

def summarize_breaches_with_llm(breaches: list, sla_info: dict) -> str:
    if not breaches: return "No breaches detected."
    
    curr = sla_info.get("currency", "INR")
    total_formatted = f"{curr} {len(breaches) * (sla_info.get('penalty_per_breach') or 0):,}"
    worst_3 = sorted(breaches, key=lambda b: b["overage_minutes"], reverse=True)[:3]
    
    # THE FIX: Explicit, unambiguous text generation per flight
    worst_txt = "\n".join([f"- Flight {b['flight_number']}: Total Duration was {b['actual_minutes']} mins (which is {b['overage_minutes']} mins over the limit). Explicit Penalty: {curr} {b['penalty']:,}" for b in worst_3])
    
    llm = ChatGroq(model_name="llama-3.1-8b-instant", temperature=0)
    prompt = ChatPromptTemplate.from_template("""
    Write a summary of the SLA breaches using EXACTLY the numbers provided. DO NO MATH.
    
    DATA PROVIDED:
    Total Breaches: {total_breaches} | Total Penalty: {total_penalty_formatted}
    Worst Offenders: 
    {worst_offenders}
    
    CRITICAL INSTRUCTIONS:
    1. Sentence 1: State the total breaches and total penalty.
    2. Remaining Sentences: You must write exactly one distinct sentence for EACH of the worst flights listed above, using this exact template:
    "Flight [Flight Number] had a total duration of [Duration], exceeding the limit by [Overage], with an explicit penalty of [Penalty]."
    
    Do not add any extra commentary. Do NOT group flights together.
    """)
    return (prompt | llm).invoke({
        "total_breaches": len(breaches), "total_penalty_formatted": total_formatted, "worst_offenders": worst_txt
    }).content.strip()

def detect_breaches(file_bytes: bytes) -> dict:
    dfs = extract_tables_from_docx(file_bytes)
    sla_info = get_sla_thresholds(extract_text_from_docx(file_bytes))
    
    if sla_info.get("extraction_failed"): 
        return {"status": "error", "message": sla_info.get("error")}

    limit, penalty = float(sla_info.get("threshold_minutes", 45)), int(sla_info.get("penalty_per_breach") or 0)
    vendor, metric = sla_info.get("vendor"), sla_info.get("metric")

    all_breaches = []
    for df in dfs: all_breaches.extend(detect_breaches_from_df(df, limit, penalty, vendor, metric))
    all_breaches.sort(key=lambda b: b["overage_minutes"], reverse=True)

    return {
        "status": "success", "vendor": vendor, "metric": metric, "currency": sla_info.get("currency", "INR"),
        "threshold_minutes": limit, "penalty_per_breach": penalty, "total_breaches": len(all_breaches),
        "total_penalty": len(all_breaches) * penalty, "breaches": all_breaches,
        "summary": summarize_breaches_with_llm(all_breaches, sla_info)
    }